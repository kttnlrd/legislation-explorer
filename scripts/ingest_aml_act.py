#!/usr/bin/env python3
"""Ingest AML/CTF Act 2006 from DOCX into legislation-explorer format.

Single-pass approach: extract TOC, build section-content map from body.

Produces:
  - data/aml-ctf-2006/tree.json
  - data/aml-ctf-2006/sections/{part}/{division}/{section}.md

Usage:
    python3 scripts/ingest_aml_act.py [--docx PATH] [--out-dir PATH]
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from docx import Document

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
DOCX_PATH = Path.home() / "legislation-explorer" / "data" / "aml-ctf-2006" / "raw" / "aml-ctf-act-2006.docx"
OUT_DIR = Path.home() / "legislation-explorer" / "data" / "aml-ctf-2006"
SECTION_DIR = OUT_DIR / "sections"

COMPILATION_NO = 59
COMPILATION_DATE = "31 March 2025"
ACT_NAME = "Anti-Money Laundering and Counter-Terrorism Financing Act 2006"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
_EM = "\u2014"
_NBSP = "\u00a0"


def clean(text: str) -> str:
    text = text.replace(_EM, "\u2014")  # normalize em dash
    text = text.replace(_NBSP, " ")
    return text.strip()


def toc_part_id(text: str) -> str | None:
    m = re.match(r'Part\s+(\d+[A-Z]?)\s*[—\-–]', text)
    return m.group(1) if m else None


def toc_div_id(text: str) -> str | None:
    m = re.match(r'Division\s+(\d+[A-Z]?)\s*[—\-–]', text)
    return m.group(1) if m else None


def toc_subdiv_id(text: str) -> str | None:
    m = re.match(r'Subdivision\s+([A-Z])\s*[—\-–]', text)
    return m.group(1) if m else None


def extract_num(text: str) -> str | None:
    """Extract leading number from '1  Short title' or '1\tShort title\t1'."""
    m = re.match(r'(\d+[A-Z]?)(?:\s{2,}|\t)', text)
    return m.group(1) if m else None


def extract_title(text: str) -> str:
    """Extract title after leading number. Handles tabs and spaces."""
    return re.sub(r'^\d+[A-Z]?(?:\s{2,}|\t)', '', text).strip()


def strip_toc_page(text: str) -> str:
    """Remove trailing tab+page number from TOC entries."""
    return re.sub(r'\t\d+$', '', text).strip()


# ---------------------------------------------------------------------------
# Parse TOC
# ---------------------------------------------------------------------------
def parse_toc(doc: Document) -> list[dict]:
    """Parse the TOC section (paragraphs with toc styles) into a hierarchy."""
    parts = []
    current_part = None
    current_div = None
    current_subdiv = None
    in_toc = False

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style else ''

        # Detect start of TOC
        if not in_toc:
            if text == 'Contents':
                in_toc = True
            continue

        # End of TOC — body starts
        if style in ('LongT', 'ActHead'):
            break

        if style == 'toc 1':
            # Schedule 1
            m = re.match(r'Schedule\s+(\d+)\s*[—\-–]', text)
            if m:
                sch_id = m.group(1)
                current_part = {
                    'id': f'schedule-{sch_id}',
                    'title': clean(strip_toc_page(text)),
                    'divisions': [{
                        'id': '1',
                        'title': clean(strip_toc_page(text)),
                        'subdivisions': [],
                        'sections': []
                    }]
                }
                parts.append(current_part)
                current_div = current_part['divisions'][0]
                current_subdiv = None
            continue

        if style == 'toc 2' and text.startswith('Part'):
            pid = toc_part_id(text)
            if pid:
                current_part = {
                    'id': pid,
                    'title': clean(strip_toc_page(text)),
                    'divisions': []
                }
                parts.append(current_part)
                current_div = None
                current_subdiv = None
            continue

        if style == 'toc 2' and text.startswith('Endnotes'):
            break

        if style == 'toc 3':
            if current_part is None:
                continue
            did = toc_div_id(text)
            if did:
                current_div = {
                    'id': did,
                    'title': clean(strip_toc_page(text)),
                    'subdivisions': [],
                    'sections': []
                }
                current_part['divisions'].append(current_div)
                current_subdiv = None
            continue

        if style == 'toc 4':
            if current_div is None:
                continue
            sd = toc_subdiv_id(text)
            if sd:
                current_subdiv = {
                    'id': sd,
                    'title': clean(strip_toc_page(text)),
                    'sections': []
                }
                current_div['subdivisions'].append(current_subdiv)
            continue

        if style == 'toc 5':
            num = extract_num(text)
            if num:
                title = clean(extract_title(text))
                title = strip_toc_page(title)
                entry = {'id': num, 'title': title}
                if current_subdiv is not None:
                    current_subdiv['sections'].append(entry)
                elif current_div is not None:
                    current_div['sections'].append(entry)
                elif current_part is not None:
                    # Part without divisions — create implicit div
                    if not current_part['divisions']:
                        current_div = {
                            'id': current_part['id'],
                            'title': current_part['title'],
                            'subdivisions': [],
                            'sections': []
                        }
                        current_part['divisions'].append(current_div)
                    current_div['sections'].append(entry)
            continue

    return parts


# ---------------------------------------------------------------------------
# Build body map: section_id -> (title, content_text)
# ---------------------------------------------------------------------------
def build_body_map(doc: Document) -> dict[str, dict]:
    """Single pass through body paragraphs to extract per-section content.

    Uses the paragraph number as an additional key to handle duplicate
    section numbers (e.g. section 1 in Part 1 vs section 1 in Schedule 1).
    Returns dict: section_id -> {'title': str, 'content': str, 'para_num': int}
    """
    body_map = {}
    current_section_id = None
    current_title = None
    current_para = None
    content_lines = []
    in_body = False

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ''

        # Detect start of body
        if style == 'LongT' and 'An Act to combat' in text:
            in_body = True
            continue

        if not in_body:
            continue

        if style == 'ActHead 5':
            # Flush previous section
            if current_section_id is not None and current_para is not None:
                key = (current_section_id, current_para)
                body_map[key] = {
                    'title': current_title,
                    'content': _flush_content(content_lines)
                }
                content_lines = []

            num = extract_num(text)
            title = extract_title(text) if num else text
            if num:
                current_section_id = num
                current_title = clean(title)
                current_para = i
            else:
                current_section_id = None
                current_title = None
                current_para = None
            continue

        # If we hit a new part/division/subdivision heading, keep current section
        if style in ('ActHead 2', 'ActHead 3', 'ActHead 4', 'ActHead 1'):
            continue

        # Collect content for current section
        if current_section_id is not None:
            content_lines.append((style, p.text))

    # Flush last section
    if current_section_id is not None and current_para is not None:
        key = (current_section_id, current_para)
        body_map[key] = {
            'title': current_title,
            'content': _flush_content(content_lines)
        }

    return body_map


def _flush_content(lines: list) -> str:
    """Convert raw (style, text) lines into markdown."""
    md_lines = []
    for style, text in lines:
        t = text.strip()
        if not t:
            if md_lines and md_lines[-1] != '':
                md_lines.append('')
            continue

        cleaned = clean(t)

        if style in ('subsection', 'subsection2'):
            md_lines.append(cleaned)
        elif style == 'paragraph':
            md_lines.append(cleaned)
        elif style == 'paragraph(sub)':
            md_lines.append(cleaned)
        elif style == 'Definition':
            md_lines.append(f'**{cleaned}**')
        elif style in ('note(text)', 'note(para)', 'note(margin)'):
            md_lines.append(f'  > {cleaned}')
        elif style in ('BoxList', 'BoxPara'):
            md_lines.append(cleaned)
        elif style == 'Penalty':
            md_lines.append(f'Penalty: {cleaned}')
        elif style in ('SO Text', 'SO Bullet', 'SO Para'):
            md_lines.append(cleaned)
        elif style == 'SubsectionHead':
            md_lines.append(f'**{cleaned}**')
        else:
            md_lines.append(cleaned)

    # Collapse consecutive blanks
    result = []
    prev_blank = False
    for line in md_lines:
        if line == '':
            if prev_blank:
                continue
            prev_blank = True
        else:
            prev_blank = False
        result.append(line)

    # Trim leading/trailing blanks
    while result and result[0] == '':
        result.pop(0)
    while result and result[-1] == '':
        result.pop()

    return '\n'.join(result)


# ---------------------------------------------------------------------------
# Get structure titles from body, scoped by part
# ---------------------------------------------------------------------------
def get_structure_titles(doc: Document) -> tuple[dict, dict, dict]:
    """Extract part/division/subdivision titles from body headings.

    Returns dicts keyed by (part_id, div_id) or (part_id, div_id, subdiv_id)
    so titles from different parts don't collide.
    """
    part_titles = {}
    div_titles = {}  # keyed by (part_id, div_id)
    subdiv_titles = {}  # keyed by (part_id, div_id, subdiv_id)

    current_part = None
    current_div = None

    for p in doc.paragraphs:
        text = p.text.strip()
        style = p.style.name if p.style else ''

        if style == 'ActHead 2':
            m = re.match(r'Part\s+(\d+[A-Z]?)\s*[—\-–]\s*(.*)', text)
            if m:
                pid = m.group(1)
                rest = m.group(2).strip()
                t = f"Part {pid} — {rest}" if rest else f"Part {pid}"
                part_titles[pid] = clean(t)
                current_part = pid
                current_div = None
            else:
                # Schedule heading
                m2 = re.match(r'Schedule\s+(\d+)\s*[—\-–]\s*(.*)', text)
                if m2:
                    current_part = f'schedule-{m2.group(1)}'
                    current_div = None

        elif style == 'ActHead 3':
            m = re.match(r'Division\s+(\d+[A-Z]?)\s*[—\-–]\s*(.*)', text)
            if m and current_part is not None:
                did = m.group(1)
                rest = m.group(2).strip()
                t = f"Division {did} — {rest}" if rest else f"Division {did}"
                div_titles[(current_part, did)] = clean(t)
                current_div = did

        elif style == 'ActHead 4':
            if current_part is not None and current_div is not None:
                m = re.match(r'Subdivision\s+([A-Z])\s*[—\-–]\s*(.*)', text)
                if m:
                    sd = m.group(1)
                    rest = m.group(2).strip()
                    t = f"Subdivision {sd} — {rest}" if rest else f"Subdivision {sd}"
                    subdiv_titles[(current_part, current_div, sd)] = clean(t)

    return part_titles, div_titles, subdiv_titles


# ---------------------------------------------------------------------------
# Write section file
# ---------------------------------------------------------------------------
def write_section(sec_dir: Path, sec_id: str, sec_title: str, body_text: str,
                  part_id: str, part_title: str, div_id: str, div_title: str,
                  subdiv_id: str | None = None, subdiv_title: str | None = None):
    sec_dir.mkdir(parents=True, exist_ok=True)
    path = sec_dir / f"{sec_id}.md"

    with open(path, 'w', encoding='utf-8') as f:
        f.write("---\n")
        f.write(f"act: \"{ACT_NAME}\"\n")
        f.write(f"part: \"{part_id}\"\n")
        f.write(f"part_title: \"{part_title}\"\n")
        f.write(f"division: \"{div_id}\"\n")
        f.write(f"division_title: \"{div_title}\"\n")
        if subdiv_id is not None:
            f.write(f"subdivision: \"{subdiv_id}\"\n")
            f.write(f"subdivision_title: \"{subdiv_title}\"\n")
        f.write(f"section: \"{sec_id}\"\n")
        f.write(f"section_title: \"{sec_title}\"\n")
        f.write(f"compilation_no: {COMPILATION_NO}\n")
        f.write(f"compilation_date: \"{COMPILATION_DATE}\"\n")
        f.write("---\n")
        f.write("\n")
        f.write(f"# {sec_id} {sec_title}\n")
        f.write("\n")
        if body_text:
            f.write(body_text)
            f.write("\n")
        f.write("\n")
        f.write("---\n")
        f.write(f"*{ACT_NAME}*\n")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(description="Ingest AML/CTF Act 2006 from DOCX")
    parser.add_argument("--docx", type=str, default=str(DOCX_PATH))
    parser.add_argument("--out-dir", type=str, default=str(OUT_DIR))
    args = parser.parse_args()

    docx_path = Path(args.docx)
    out_dir = Path(args.out_dir)
    section_dir = out_dir / "sections"

    print(f"Reading DOCX: {docx_path}")
    doc = Document(str(docx_path))

    # 1. Parse TOC
    print("Parsing TOC...")
    parts = parse_toc(doc)
    print(f"Found {len(parts)} parts")

    # 2. Get structure titles from body (scoped by part)
    part_titles, div_titles, subdiv_titles = get_structure_titles(doc)

    # 3. Build section body map (single pass)
    print("Building section content map...")
    body_map = build_body_map(doc)
    # body_map keys are (section_id, para_num) tuples
    print(f"Found {len(body_map)} body section entries")

    # 4. Collect all TOC section IDs to know which are real sections
    toc_section_ids = set()
    for part in parts:
        for div in part['divisions']:
            for s in div['sections']:
                toc_section_ids.add(s['id'])
            for sd in div['subdivisions']:
                for s in sd['sections']:
                    toc_section_ids.add(s['id'])

    # 5. Build tree & write section files
    print("Building tree.json and writing section files...")
    tree = {
        "act": ACT_NAME,
        "compilation_no": COMPILATION_NO,
        "compilation_date": COMPILATION_DATE,
        "parts": []
    }

    total_sections = 0

    for part in parts:
        pid = part['id']
        # Use body title if available, else TOC title
        pt = clean(part_titles.get(pid, part['title']))

        part_entry = {
            "id": pid,
            "title": pt,
            "divisions": []
        }

        for div in part['divisions']:
            did = div['id']
            # Use body title scoped by part+division
            dt = clean(div_titles.get((pid, did), div['title']))

            div_entry = {
                "id": did,
                "title": dt,
                "subdivisions": [],
                "sections": []
            }

            # Sections directly under division
            for sec in div['sections']:
                sid = sec['id']
                # Find body content — use the first occurrence of this section ID
                # that is closest to the part's position in the document
                body_info = _find_body_entry(body_map, sid, pid, doc)
                st = body_info.get('title') or clean(sec['title'])
                body_text = body_info.get('content', '')
                sec_path = f"part-{pid}/division-{did}/{sid}.md"

                div_entry['sections'].append({
                    "id": sid,
                    "title": st,
                    "path": sec_path
                })

                sec_dir = section_dir / f"part-{pid}" / f"division-{did}"
                write_section(sec_dir, sid, st, body_text, pid, pt, did, dt)
                total_sections += 1

            # Subdivisions
            for subdiv in div['subdivisions']:
                sd_id = subdiv['id']
                sd_title = clean(subdiv_titles.get((pid, did, sd_id), subdiv['title']))

                subdiv_entry = {
                    "id": sd_id,
                    "title": sd_title,
                    "sections": []
                }

                for sec in subdiv['sections']:
                    sid = sec['id']
                    body_info = _find_body_entry(body_map, sid, pid, doc)
                    st = body_info.get('title') or clean(sec['title'])
                    body_text = body_info.get('content', '')
                    sec_path = f"part-{pid}/division-{did}/{sid}.md"

                    subdiv_entry['sections'].append({
                        "id": sid,
                        "title": st,
                        "path": sec_path
                    })

                    sec_dir = section_dir / f"part-{pid}" / f"division-{did}"
                    write_section(sec_dir, sid, st, body_text, pid, pt, did, dt,
                                  sd_id, sd_title)
                    total_sections += 1

                div_entry['subdivisions'].append(subdiv_entry)

            part_entry['divisions'].append(div_entry)

        tree['parts'].append(part_entry)

    # Write tree.json
    tree_path = out_dir / "tree.json"
    with open(tree_path, 'w', encoding='utf-8') as f:
        json.dump(tree, f, indent=2, ensure_ascii=False)
    print(f"Written: {tree_path}")

    # Summary
    sec_files = list(section_dir.rglob("*.md"))
    print(f"\n{'='*60}")
    print(f"SUMMARY")
    print(f"{'='*60}")
    print(f"Parts: {len(parts)}")
    print(f"Divisions: {sum(len(p['divisions']) for p in parts)}")
    print(f"Subdivisions: {sum(len(d['subdivisions']) for p in parts for d in p['divisions'])}")
    print(f"Section files written: {total_sections}")
    print(f"Section files on disk: {len(sec_files)}")
    print(f"Tree file: {tree_path}")
    print(f"Done.")


def _find_body_entry(body_map: dict, section_id: str, part_id: str, doc) -> dict:
    """Find the best body entry for a section, handling duplicate section numbers.

    Uses para_num proximity to the part heading to pick the right one.
    """
    candidates = [(k, v) for k, v in body_map.items() if k[0] == section_id]
    if not candidates:
        return {}
    if len(candidates) == 1:
        return candidates[0][1]

    # Multiple candidates — find the part/schedule heading para number
    target_para = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ''
        if style == 'ActHead 2':
            m = re.match(r'Part\s+' + re.escape(part_id) + r'\s*[—\-–]', text)
            if m:
                target_para = i
                break
        # Handle Schedule headings
        if style == 'ActHead 1' and part_id.startswith('schedule-'):
            m = re.match(r'Schedule\s+' + re.escape(part_id.replace('schedule-', '')) + r'\s*[—\-–]', text)
            if m:
                target_para = i
                break

    if target_para is None:
        return candidates[0][1]

    # Find the candidate whose para_num is closest to (but after) the part heading
    best = None
    best_dist = float('inf')
    for (sid, pnum), v in candidates:
        if pnum >= target_para:
            dist = pnum - target_para
            if dist < best_dist:
                best_dist = dist
                best = v

    return best or candidates[0][1]


if __name__ == "__main__":
    main()