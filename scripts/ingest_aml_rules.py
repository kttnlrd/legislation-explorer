#!/usr/bin/env python3
"""
Ingest AML/CTF Rules 2007 DOCX → tree.json + section markdown files.

Structure: Chapter → Part → Section
Chapters 1–10 have Parts (e.g. Part 1.1, 1.2, … 10.3).
Chapters 11+ have no parts – sections sit directly under the chapter.

Section numbering:
  "1 Name of Instrument", "2 Rules" — preliminary sections before any chapter.
  "1.1.1", "1.2.1", …, "10.3.X" — sections under a chapter-part.
  "11.1", "11.2", …, "81.X" — sections under a chapter without parts.
"""

import json
import os
import re
import sys

from docx import Document

DOCX_PATH = "data/aml-ctf-2006/raw/aml-ctf-rules-2007.docx"
OUT_DIR = "data/aml-ctf-rules-2007"
OUT_TREE = os.path.join(OUT_DIR, "tree.json")
OUT_SECTIONS = os.path.join(OUT_DIR, "sections")

ACT_TITLE = "Anti-Money Laundering and Counter-Terrorism Financing Rules Instrument 2007 (No. 1)"
COMPILATION_NO = 76
COMPILATION_DATE = "22 November 2024"

# ── helpers ──────────────────────────────────────────────────────────

def slugify(text):
    """Make a filesystem-safe slug."""
    s = text.strip().lower()
    s = re.sub(r"[^a-z0-9]+", "-", s)
    s = s.strip("-")
    return s or "unnamed"


def chapter_id_from_num(n):
    return f"ch-{n}"


def part_id_from_num(n):
    # E.g. "1.1" → "part-1.1"
    return f"part-{n}"


def section_id_from_num(n):
    return str(n)


# ── parsing ──────────────────────────────────────────────────────────

def parse_docx(path):
    """Return list of ordered structure items: chapters, parts, sections."""
    doc = Document(path)
    pars = doc.paragraphs

    # Step 1: locate the body start (after TOC) and end (before endnotes)
    body_start = 0
    body_end = len(pars)
    for i, p in enumerate(pars):
        style = p.style.name if p.style else ""
        text = p.text.strip()
        if style == "HR" and text.startswith("1\tName of Instrument"):
            body_start = i
            break
    for i in range(len(pars) - 1, -1, -1):
        p = pars[i]
        style = p.style.name if p.style else ""
        if style == "Note Heading":
            body_end = i
            break

    print(f"Body paragraphs: {body_start} – {body_end}")

    # Step 2: extract ordered structure (chapter → part → sections)
    # We'll walk paragraph-by-paragraph tracking current chapter/part

    items = []  # list of dicts: {type: 'chapter'|'part'|'section', ...}
    current_chapter = None
    current_part = None

    # section_number → section_title mapping for body text
    section_map = {}  # {section_num: title_text}
    section_order = []  # ordered list of section numbers as encountered

    # Regex for section heads like "1.1.1 ...", "11.1 ...", "1 Name of Instrument"
    sec_head_re = re.compile(r"^(\d+(?:\.\d+)*)\s+(.*)")

    # Patterns for detecting chapter/part/section at each paragraph
    # We also need to handle "Part 3.2" etc which aren't HD styled

    # First pass: find chapter/part boundaries and section headings
    i = body_start
    while i < body_end:
        p = pars[i]
        style = p.style.name if p.style else ""
        text = p.text.strip()
        if not text:
            i += 1
            continue

        # ── Chapter heading (HP style) ──
        if style == "HP":
            # e.g. "CHAPTER 1" or "CHAPTER 3\tCorrespondent Banking"
            m = re.match(r"CHAPTER\s+(\d+)(?:\s+(.*))?", text)
            if m:
                ch_num = int(m.group(1))
                ch_title = (m.group(2) or "").strip()
                current_chapter = {
                    "num": ch_num,
                    "title": f"Ch {ch_num}" + (f" — {ch_title}" if ch_title else ""),
                    "title_raw": ch_title,
                }
                current_part = None
                items.append({"type": "chapter", "data": current_chapter})
            i += 1
            continue

        # ── Part heading (HD style, or "Part X.Y" pattern in text) ──
        if style == "HD" or re.match(r"^Part\s+\d+\.\d+\s", text):
            # e.g. "Part 1.1\tIntroduction" or "Part 3.2 \tOngoing assessments..."
            m = re.match(r"Part\s+(\d+\.\d+)\s+(.*)", text)
            if m:
                part_num = m.group(1)
                part_title = m.group(2).strip()
                current_part = {
                    "num": part_num,
                    "title": f"Part {part_num} — {part_title}",
                    "title_raw": part_title,
                }
                items.append({"type": "part", "data": current_part})
            i += 1
            continue

        # ── Section heading ──
        # Sections can be:
        #   - "1 Name of Instrument" (preliminary, before any chapter)
        #   - "2 Rules" (preliminary)
        #   - "1.1.1 These Rules..." (chapter-part sections)
        #   - "11.1 These Rules..." (chapter sections without parts)
        m = sec_head_re.match(text)
        if m:
            sec_num = m.group(1)
            sec_title = m.group(2).strip() if m.group(2) else ""
            # Skip if this is a short number that's not a section heading
            # (e.g. "(1)" is not a section heading)
            parts = sec_num.split(".")
            if len(parts) <= 3:
                # Check it's really a section heading, not a sub-paragraph reference
                # Section numbers are:
                #   - Simple digits: 1, 2 (preliminary)
                #   - C.P.S: 1.1.1, 4.2.14, etc.
                #   - Chapter.S: 11.1, 12.1, etc.
                is_valid = False
                if len(parts) == 1:
                    # Must be "1" or "2" in preliminary area (before any chapter)
                    if sec_num in ("1", "2") and current_chapter is None:
                        is_valid = True
                elif len(parts) == 2:
                    # Chapter.Section format - chapters without parts (11+)
                    # But must NOT be a cross-reference like "4.2" in body text
                    # Check that this looks like a section heading, not a reference
                    ch = int(parts[0])
                    if ch >= 11:
                        is_valid = True
                elif len(parts) == 3:
                    # C.P.S format - need to check it matches current chapter/part
                    ch = int(parts[0])
                    # Verify it belongs to a chapter (not a random 3-part number)
                    if ch >= 1:
                        is_valid = True

                if is_valid and sec_title:
                    # This is a section heading
                    section_order.append(sec_num)
                    section_map[sec_num] = sec_title
                    items.append({
                        "type": "section",
                        "data": {
                            "num": sec_num,
                            "title": sec_title,
                            "para_start": i,
                        }
                    })

        i += 1

    print(f"Found {len(items)} structure items")
    chs = sum(1 for it in items if it["type"] == "chapter")
    pts = sum(1 for it in items if it["type"] == "part")
    secs = sum(1 for it in items if it["type"] == "section")
    print(f"  Chapters: {chs}, Parts: {pts}, Sections: {secs}")

    return items, doc, body_start, body_end


# ── extract section content ──────────────────────────────────────────

def extract_section_content(doc, body_start, body_end, items):
    """For each section item, find its end and extract the full text content."""
    section_contents = {}

    # Find section items
    section_items = [(i, it) for i, it in enumerate(items) if it["type"] == "section"]

    for idx, (item_idx, item) in enumerate(section_items):
        sec_num = item["data"]["num"]
        start_para = item["data"]["para_start"]

        # Determine end paragraph: next section, part, or chapter
        end_para = body_end
        for j in range(item_idx + 1, len(items)):
            next_item = items[j]
            if next_item["type"] in ("section", "part", "chapter"):
                end_para = next_item["data"].get("para_start", body_end)
                break

        # Extract text from start_para to end_para
        lines = []
        # Include the heading paragraph
        heading_text = doc.paragraphs[start_para].text.strip()
        lines.append(heading_text)

        for pi in range(start_para + 1, end_para):
            p = doc.paragraphs[pi]
            style = p.style.name if p.style else ""
            text = p.text

            # Skip empty paragraphs and Note headings before endnotes
            if not text.strip():
                continue
            if style == "Note Heading":
                continue

            lines.append(text)

        # Join and clean up
        content = "\n".join(lines)
        section_contents[sec_num] = content

    return section_contents


# ── build tree ────────────────────────────────────────────────────────

def build_tree(items, section_contents):
    """Build the tree.json structure."""
    tree = {
        "act": ACT_TITLE,
        "compilation_no": COMPILATION_NO,
        "compilation_date": COMPILATION_DATE,
        "parts": [],
    }

    current_chapter = None
    current_part = None  # dict with "id", "title", "sections"

    # Preliminary sections (before chapter 1)
    prelim_chapter = {
        "id": "preliminary",
        "title": "Preliminary",
        "divisions": [],
        "sections": [],
    }

    for item in items:
        if item["type"] == "chapter":
            # Save previous part if any
            if current_part is not None and current_chapter is not None:
                current_chapter["parts"].append(current_part)
                current_part = None

            # Save previous chapter
            if current_chapter is not None:
                tree["parts"].append(current_chapter)

            ch = item["data"]
            current_chapter = {
                "id": chapter_id_from_num(ch["num"]),
                "title": ch["title"],
                "parts": [],
                "sections": [],  # for chapters without parts
            }

        elif item["type"] == "part":
            # Save previous part
            if current_part is not None and current_chapter is not None:
                current_chapter["parts"].append(current_part)

            pt = item["data"]
            current_part = {
                "id": part_id_from_num(pt["num"]),
                "title": pt["title"],
                "sections": [],
            }

        elif item["type"] == "section":
            sec = item["data"]
            sec_num = sec["num"]
            sec_title = sec["title"]

            # Determine where this section belongs
            path_parts = []
            section_entry = {
                "id": section_id_from_num(sec_num),
                "title": sec_title,
            }

            # Check if it's a preliminary section (1, 2)
            parts_list = sec_num.split(".")
            if len(parts_list) == 1 and current_chapter is None:
                # Preliminary section
                prelim_chapter["sections"].append(section_entry)
                continue

            # Check if we're in a chapter with parts
            if current_part is not None:
                current_part["sections"].append(section_entry)
            elif current_chapter is not None:
                current_chapter["sections"].append(section_entry)
            else:
                # Fallback: preliminary
                prelim_chapter["sections"].append(section_entry)

    # Flush last part/chapter
    if current_part is not None and current_chapter is not None:
        current_chapter["parts"].append(current_part)
    if current_chapter is not None:
        tree["parts"].append(current_chapter)

    # Add preliminary sections if any
    if prelim_chapter["sections"]:
        # Insert at beginning
        tree["parts"].insert(0, prelim_chapter)

    return tree


# ── determine file path for a section ────────────────────────────────

def get_section_path(sec_num, tree, items):
    """Determine the correct file path for a section."""
    parts_list = sec_num.split(".")

    if len(parts_list) == 1:
        # Preliminary section (1, 2)
        return f"preliminary/{sec_num}.md"

    ch = int(parts_list[0])

    if len(parts_list) == 2:
        # Chapter section without parts
        return f"ch-{ch}/{sec_num}.md"

    if len(parts_list) == 3:
        # Chapter.Part.Section
        part_num = f"{parts_list[0]}.{parts_list[1]}"
        return f"ch-{ch}/part-{part_num}/{sec_num}.md"

    return f"misc/{sec_num}.md"


def get_path_in_tree(sec_num, tree):
    """Get the path string stored in tree.json for a section."""
    parts_list = sec_num.split(".")

    if len(parts_list) == 1:
        return f"preliminary/{sec_num}.md"

    ch = int(parts_list[0])

    if len(parts_list) == 2:
        return f"ch-{ch}/{sec_num}.md"

    if len(parts_list) == 3:
        part_num = f"{parts_list[0]}.{parts_list[1]}"
        return f"ch-{ch}/part-{part_num}/{sec_num}.md"

    return f"misc/{sec_num}.md"


# ── write section markdown files ─────────────────────────────────────

def write_section_files(section_contents, tree, items):
    """Write each section's markdown file with YAML frontmatter."""
    for sec_num, content in section_contents.items():
        # Determine chapter/part info from section number
        parts_list = sec_num.split(".")
        ch_num = None
        part_num = None
        ch_title = ""
        part_title = ""

        if len(parts_list) == 1:
            # Preliminary
            ch_num = 0
            ch_title = "Preliminary"
            part_num = ""
            part_title = ""
        elif len(parts_list) == 2:
            ch_num = int(parts_list[0])
        elif len(parts_list) == 3:
            ch_num = int(parts_list[0])
            part_num = f"{parts_list[0]}.{parts_list[1]}"

        # Get chapter title from tree
        if ch_num and ch_num > 0:
            for p in tree["parts"]:
                if p["id"] == chapter_id_from_num(ch_num):
                    ch_title = p["title"]
                    break

        # Get part title
        if part_num and ch_num and ch_num > 0:
            for p in tree["parts"]:
                if p["id"] == chapter_id_from_num(ch_num):
                    for pt in p.get("parts", []):
                        if pt["id"] == part_id_from_num(part_num):
                            part_title = pt["title"]
                            break
                    break

        # Get section title
        sec_title = ""
        parts_list_sec = sec_num.split(".")
        if len(parts_list_sec) == 1:
            for p in tree["parts"]:
                if p["id"] == "preliminary":
                    for s in p.get("sections", []):
                        if s["id"] == sec_num:
                            sec_title = s["title"]
                            break
                    break
        elif len(parts_list_sec) == 2:
            ch_id = chapter_id_from_num(int(parts_list_sec[0]))
            for p in tree["parts"]:
                if p["id"] == ch_id:
                    for s in p.get("sections", []):
                        if s["id"] == sec_num:
                            sec_title = s["title"]
                            break
                    break
        elif len(parts_list_sec) == 3:
            ch_id = chapter_id_from_num(int(parts_list_sec[0]))
            part_id = part_id_from_num(f"{parts_list_sec[0]}.{parts_list_sec[1]}")
            for p in tree["parts"]:
                if p["id"] == ch_id:
                    for pt in p.get("parts", []):
                        if pt["id"] == part_id:
                            for s in pt.get("sections", []):
                                if s["id"] == sec_num:
                                    sec_title = s["title"]
                                    break
                    break

        # Build file path
        rel_path = get_path_in_tree(sec_num, tree)
        abs_path = os.path.join(OUT_SECTIONS, rel_path)

        os.makedirs(os.path.dirname(abs_path), exist_ok=True)

        # Extract first line (the heading) from content
        first_line = content.split("\n")[0] if content else ""
        # Remove the section number prefix for the markdown heading
        heading_match = re.match(rf"^{re.escape(sec_num)}\s+(.*)", first_line)
        display_title = heading_match.group(1) if heading_match else sec_title

        # Build YAML frontmatter
        yaml_lines = ["---"]
        yaml_lines.append(f'act: "{ACT_TITLE}"')
        if ch_num and ch_num > 0:
            yaml_lines.append(f'chapter: "{ch_title}"')
        else:
            yaml_lines.append(f'chapter: ""')
        if part_num:
            yaml_lines.append(f'part: "{part_num}"')
            yaml_lines.append(f'part_title: "{part_title}"')
        else:
            yaml_lines.append(f'part: ""')
            yaml_lines.append(f'part_title: ""')
        yaml_lines.append(f'section: "{sec_num}"')
        yaml_lines.append(f'section_title: "{sec_title}"')
        yaml_lines.append(f"compilation_no: {COMPILATION_NO}")
        yaml_lines.append(f'compilation_date: "{COMPILATION_DATE}"')
        yaml_lines.append(f'source_docx: "aml-ctf-rules-2007"')
        yaml_lines.append("---")
        yaml_lines.append("")

        # Write the markdown section heading
        yaml_lines.append(f"# {sec_num} {display_title}")
        yaml_lines.append("")

        # Write body content (skip the first line which is the heading)
        body_lines = content.split("\n")[1:]
        # Remove empty leading lines
        while body_lines and not body_lines[0].strip():
            body_lines = body_lines.pop(0) if False else body_lines
        # But we need a better approach
        non_empty_start = 0
        for j, line in enumerate(body_lines):
            if line.strip():
                non_empty_start = j
                break
        body_lines = body_lines[non_empty_start:]

        for line in body_lines:
            yaml_lines.append(line)

        # Write footer
        yaml_lines.append("")
        yaml_lines.append("---")
        yaml_lines.append(f"*{ACT_TITLE}*")

        with open(abs_path, "w", encoding="utf-8") as f:
            f.write("\n".join(yaml_lines))

        # Update tree.json path
        # (We'll do this after building the tree)


def update_tree_paths(tree, section_contents):
    """Update tree.json with correct file paths."""
    for sec_num in section_contents:
        parts_list = sec_num.split(".")
        path = get_path_in_tree(sec_num, tree)

        if len(parts_list) == 1:
            for p in tree["parts"]:
                if p["id"] == "preliminary":
                    for s in p.get("sections", []):
                        if s["id"] == sec_num:
                            s["path"] = path
                            break
                    break
        elif len(parts_list) == 2:
            ch_id = chapter_id_from_num(int(parts_list[0]))
            for p in tree["parts"]:
                if p["id"] == ch_id:
                    for s in p.get("sections", []):
                        if s["id"] == sec_num:
                            s["path"] = path
                            break
                    break
        elif len(parts_list) == 3:
            ch_id = chapter_id_from_num(int(parts_list[0]))
            part_id = part_id_from_num(f"{parts_list[0]}.{parts_list[1]}")
            for p in tree["parts"]:
                if p["id"] == ch_id:
                    for pt in p.get("parts", []):
                        if pt["id"] == part_id:
                            for s in pt.get("sections", []):
                                if s["id"] == sec_num:
                                    s["path"] = path
                                    break
                    break


# ── main ──────────────────────────────────────────────────────────────

def main():
    docx_path = os.path.join(os.getcwd(), DOCX_PATH)
    if not os.path.exists(docx_path):
        print(f"ERROR: DOCX not found at {docx_path}")
        sys.exit(1)

    print(f"Parsing: {docx_path}")
    items, doc, body_start, body_end = parse_docx(docx_path)

    print("\nExtracting section content...")
    section_contents = extract_section_content(doc, body_start, body_end, items)
    print(f"  Extracted {len(section_contents)} section contents")

    print("\nBuilding tree...")
    tree = build_tree(items, section_contents)
    print(f"  Tree has {len(tree['parts'])} chapter/part entries")

    print("\nWriting section files...")
    os.makedirs(OUT_SECTIONS, exist_ok=True)
    write_section_files(section_contents, tree, items)

    print("\nUpdating tree paths...")
    update_tree_paths(tree, section_contents)

    print("\nWriting tree.json...")
    with open(OUT_TREE, "w", encoding="utf-8") as f:
        json.dump(tree, f, indent=2, ensure_ascii=False)
    print(f"  Written: {OUT_TREE}")

    # Summary
    total_sections = sum(
        1 for p in tree["parts"]
        for s in p.get("sections", [])
    )
    for p in tree["parts"]:
        for pt in p.get("parts", []):
            total_sections += len(pt.get("sections", []))

    print(f"\n{'='*60}")
    print(f"COMPLETE: {len(tree['parts'])} chapters, {total_sections} sections")
    print(f"Tree: {OUT_TREE}")
    print(f"Sections: {OUT_SECTIONS}")
    print(f"{'='*60}")

    # Print chapter structure
    for p in tree["parts"]:
        ch_title = p["title"][:80]
        sec_count = len(p.get("sections", []))
        parts_count = len(p.get("parts", []))
        if parts_count > 0:
            print(f"  {p['id']}: {ch_title} ({parts_count} parts)")
            for pt in p["parts"]:
                print(f"    {pt['id']}: {pt['title'][:70]} ({len(pt['sections'])} sections)")
        else:
            print(f"  {p['id']}: {ch_title} ({sec_count} sections)")


if __name__ == "__main__":
    main()